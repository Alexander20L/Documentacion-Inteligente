from __future__ import annotations

import io
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from .canonical import stable_hash
from .models import (
    ArtifactFormat,
    C4Element,
    CanonicalC4Model,
    ElementKind,
    RenderedArtifact,
)
from .validation import assert_valid_c4_model


def _dsl_text(value: str | None) -> str:
    return (value or "").replace("\\", "\\\\").replace('"', '\\"').replace("\r", "").replace("\n", "\\n")


def _alias(identifier: str) -> str:
    readable = re.sub(r"[^A-Za-z0-9_]", "_", identifier).strip("_") or "element"
    if readable[0].isdigit():
        readable = f"e_{readable}"
    return f"{readable}_{stable_hash(identifier)[:8]}"


def _tags(element: C4Element) -> tuple[str, ...]:
    tags = set(element.tags)
    if element.kind == ElementKind.EXTERNAL_SYSTEM:
        tags.add("External")
    if element.provenance.value == "analyst_provided":
        tags.add("AnalystProvided")
    if element.provenance.value == "inferred":
        tags.add("ApprovedInference")
    return tuple(sorted(tags))


def render_structurizr_dsl(model: CanonicalC4Model) -> str:
    assert_valid_c4_model(model)
    aliases = {item.id: _alias(item.id) for item in model.elements}
    children: dict[str, list[C4Element]] = {}
    for element in model.elements:
        if element.parent_id is not None:
            children.setdefault(element.parent_id, []).append(element)

    lines = [f'workspace "{_dsl_text(model.name)}" "{_dsl_text(model.description)}" {{', "    model {"]

    def emit(element: C4Element, indent: int) -> None:
        prefix = " " * indent
        alias = aliases[element.id]
        keyword = {
            ElementKind.PERSON: "person",
            ElementKind.SOFTWARE_SYSTEM: "softwareSystem",
            ElementKind.EXTERNAL_SYSTEM: "softwareSystem",
            ElementKind.CONTAINER: "container",
            ElementKind.COMPONENT: "component",
        }[element.kind]
        technology = f' "{_dsl_text(element.technology)}"' if element.kind in (ElementKind.CONTAINER, ElementKind.COMPONENT) else ""
        nested = sorted(children.get(element.id, []), key=lambda item: item.id)
        tags = _tags(element)
        has_body = bool(nested or tags)
        suffix = " {" if has_body else ""
        lines.append(f'{prefix}{alias} = {keyword} "{_dsl_text(element.name)}" "{_dsl_text(element.description)}"{technology}{suffix}')
        if has_body:
            if tags:
                lines.append(f'{prefix}    tags "{_dsl_text(",".join(tags))}"')
            for child in nested:
                emit(child, indent + 4)
            lines.append(f"{prefix}}}")

    roots = sorted((item for item in model.elements if item.parent_id is None), key=lambda item: item.id)
    for root in roots:
        emit(root, 8)
    for relation in model.relationships:
        technology = f' "{_dsl_text(relation.technology)}"' if relation.technology else ""
        lines.append(
            f'        {aliases[relation.source_id]} -> {aliases[relation.target_id]} '
            f'"{_dsl_text(relation.description)}"{technology}'
        )
    lines.extend(["    }", "", "    views {"])
    systems = sorted((item for item in model.elements if item.kind == ElementKind.SOFTWARE_SYSTEM), key=lambda item: item.id)
    for system in systems:
        lines.extend([
            f'        systemContext {aliases[system.id]} "context_{aliases[system.id]}" {{',
            "            include *",
            "            autoLayout lr",
            "        }",
            f'        container {aliases[system.id]} "containers_{aliases[system.id]}" {{',
            "            include *",
            "            autoLayout lr",
            "        }",
        ])
    containers = sorted((item for item in model.elements if item.kind == ElementKind.CONTAINER), key=lambda item: item.id)
    for container in containers:
        if children.get(container.id):
            lines.extend([
                f'        component {aliases[container.id]} "components_{aliases[container.id]}" {{',
                "            include *",
                "            autoLayout lr",
                "        }",
            ])
    lines.extend([
        "        styles {",
        '            element "External" {',
        "                background #999999",
        "                color #ffffff",
        "            }",
        '            element "AnalystProvided" {',
        "                stroke #116466",
        "                color #0b3133",
        "            }",
        '            element "ApprovedInference" {',
        "                stroke #c07b20",
        "                color #57340b",
        "            }",
        "        }",
        "    }",
        "}",
        "",
    ])
    return "\n".join(lines)


def _markdown_text(value: str | None) -> str:
    return (value or "").replace("|", "\\|").replace("\r", " ").replace("\n", " ").strip()


def render_markdown(model: CanonicalC4Model, diagram_filenames: Iterable[str] = ()) -> str:
    assert_valid_c4_model(model)
    element_by_id = {item.id: item for item in model.elements}
    diagrams = tuple(sorted(diagram_filenames))
    lines = [
        f"# {_markdown_text(model.name)}",
        "",
        _markdown_text(model.description),
        "",
        "## C4 views",
        "",
    ]
    if diagrams:
        for filename in diagrams:
            title = Path(filename).stem.replace("_", " ").replace("-", " ").title()
            lines.extend([f"### {title}", "", f"![{title}]({filename})", ""])
    else:
        lines.extend(["No rendered views were supplied.", ""])
    lines.extend([
        "## C4 elements",
        "",
        "| Level | Name | Parent | Technology | Basis | Description | Evidence |",
        "| --- | --- | --- | --- | --- | --- | --- |",
    ])
    for element in model.elements:
        parent = element_by_id[element.parent_id].name if element.parent_id else ""
        lines.append(
            f"| {element.kind.value} | {_markdown_text(element.name)} | {_markdown_text(parent)} | "
            f"{_markdown_text(element.technology)} | {element.provenance.value} | {_markdown_text(element.description)} | "
            f"{', '.join(f'`{item}`' for item in element.evidence_ids)} |"
        )
    lines.extend([
        "",
        "## Relationships",
        "",
        "| Source | Relationship | Target | Technology | Basis | Evidence |",
        "| --- | --- | --- | --- | --- | --- |",
    ])
    for relation in model.relationships:
        lines.append(
            f"| {_markdown_text(element_by_id[relation.source_id].name)} | {_markdown_text(relation.description)} | "
            f"{_markdown_text(element_by_id[relation.target_id].name)} | {_markdown_text(relation.technology)} | {relation.provenance.value} | "
            f"{', '.join(f'`{item}`' for item in relation.evidence_ids)} |"
        )
    lines.extend([
        "",
        "## Evidence index",
        "",
        "| ID | Source | Kind | Locator | SHA-256 |",
        "| --- | --- | --- | --- | --- |",
    ])
    for item in model.evidence:
        lines.append(
            f"| `{item.id}` | {item.source.value} | {item.kind.value} | {_markdown_text(item.locator)} | `{item.content_hash}` |"
        )
    lines.extend(["", "## Human decisions", ""])
    if model.decisions:
        for decision in model.decisions:
            lines.append(
                f"- `{decision.target_id}`: **{decision.decision.value}** by {_markdown_text(decision.reviewer)}. "
                f"{_markdown_text(decision.rationale)}"
            )
    else:
        lines.append("No human decisions were required for detected elements.")
    lines.extend(["", f"Canonical model: `{model.model_id}`", f"Content hash: `{model.content_hash}`", ""])
    return "\n".join(lines)


def _deterministic_docx_archive(data: bytes) -> bytes:
    source = io.BytesIO(data)
    target = io.BytesIO()
    with zipfile.ZipFile(source, "r") as input_zip, zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as output_zip:
        for name in sorted(input_zip.namelist()):
            original = input_zip.getinfo(name)
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = original.external_attr
            info.create_system = original.create_system
            output_zip.writestr(info, input_zip.read(name))
    return target.getvalue()


def render_docx(model: CanonicalC4Model, diagrams: Iterable[tuple[str, bytes]] = ()) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt

    assert_valid_c4_model(model)
    document = Document()
    fixed_time = datetime(2000, 1, 1, tzinfo=timezone.utc)
    document.core_properties.title = model.name
    document.core_properties.subject = "Approved canonical C4 model"
    document.core_properties.author = "Documentacion Inteligente"
    document.core_properties.created = fixed_time
    document.core_properties.modified = fixed_time
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10)
    document.add_heading(model.name, level=0)
    if model.description:
        document.add_paragraph(model.description)
    diagram_list = tuple(sorted(diagrams, key=lambda item: item[0]))
    document.add_heading("C4 views", level=1)
    if diagram_list:
        for name, content in diagram_list:
            document.add_heading(name, level=2)
            document.add_picture(io.BytesIO(content), width=Inches(6.5))
    else:
        document.add_paragraph("No rendered views were supplied.")
    document.add_heading("C4 elements", level=1)
    element_by_id = {item.id: item for item in model.elements}
    for element in model.elements:
        document.add_heading(element.name, level=2)
        parent = element_by_id[element.parent_id].name if element.parent_id else "None"
        document.add_paragraph(f"Type: {element.kind.value}")
        document.add_paragraph(f"Basis: {element.provenance.value}")
        document.add_paragraph(f"Parent: {parent}")
        if element.technology:
            document.add_paragraph(f"Technology: {element.technology}")
        if element.description:
            document.add_paragraph(element.description)
        document.add_paragraph(f"Evidence: {', '.join(element.evidence_ids)}")
    document.add_heading("Relationships", level=1)
    for relation in model.relationships:
        text = f"{element_by_id[relation.source_id].name} -> {element_by_id[relation.target_id].name}: {relation.description}"
        if relation.technology:
            text += f" ({relation.technology})"
        text += f" [{relation.provenance.value}]"
        document.add_paragraph(text, style="List Bullet")
        document.add_paragraph(f"Evidence: {', '.join(relation.evidence_ids)}")
    document.add_heading("Evidence index", level=1)
    for item in model.evidence:
        document.add_paragraph(
            f"{item.id}: {item.source.value}/{item.kind.value}, {item.locator}, SHA-256 {item.content_hash}",
            style="List Bullet",
        )
    document.add_heading("Human decisions", level=1)
    if model.decisions:
        for decision in model.decisions:
            document.add_paragraph(
                f"{decision.target_id}: {decision.decision.value} by {decision.reviewer}. {decision.rationale}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("No human decisions were required for detected elements.")
    document.add_paragraph(f"Canonical model: {model.model_id}")
    document.add_paragraph(f"Content hash: {model.content_hash}")
    output = io.BytesIO()
    document.save(output)
    return _deterministic_docx_archive(output.getvalue())


def render_structurizr_artifact(model: CanonicalC4Model, filename: str = "workspace.dsl") -> RenderedArtifact:
    content = render_structurizr_dsl(model)
    return RenderedArtifact(format=ArtifactFormat.STRUCTURIZR_DSL, media_type="text/plain; charset=utf-8", filename=filename, content_hash=stable_hash(content), content=content)


def render_markdown_artifact(
    model: CanonicalC4Model,
    filename: str = "ARCHITECTURE.md",
    diagram_filenames: Iterable[str] = (),
) -> RenderedArtifact:
    content = render_markdown(model, diagram_filenames)
    return RenderedArtifact(format=ArtifactFormat.MARKDOWN, media_type="text/markdown; charset=utf-8", filename=filename, content_hash=stable_hash(content), content=content)


def render_docx_artifact(
    model: CanonicalC4Model,
    filename: str = "ARCHITECTURE.docx",
    diagrams: Iterable[tuple[str, bytes]] = (),
) -> RenderedArtifact:
    content = render_docx(model, diagrams)
    return RenderedArtifact(format=ArtifactFormat.DOCX, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename, content_hash=stable_hash(content.hex()), content=content)
