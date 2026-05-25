from docx import Document
from docx.shared import Pt


def generar_word(documentacion: str, ruta_salida: str):
    documento = Document()

    estilo_normal = documento.styles["Normal"]
    estilo_normal.font.name = "Arial"
    estilo_normal.font.size = Pt(11)

    for linea in documentacion.split("\n"):
        linea = linea.strip()

        if not linea:
            continue

        if linea.startswith("# "):
            documento.add_heading(linea.replace("# ", ""), level=1)

        elif linea.startswith("## "):
            documento.add_heading(linea.replace("## ", ""), level=2)

        elif linea.startswith("### "):
            documento.add_heading(linea.replace("### ", ""), level=3)

        elif linea.startswith("- "):
            documento.add_paragraph(linea.replace("- ", ""), style="List Bullet")

        else:
            documento.add_paragraph(linea)

    documento.save(ruta_salida)